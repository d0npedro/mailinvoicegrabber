"""
Text extraction from PDF, image (OCR), and DOCX file formats.

Extraction pipeline:
  PDF  → pdfplumber (primary) → PyPDF2 (fallback)
  Image → pytesseract OCR (German + English)
  DOCX → python-docx (paragraphs + tables)
"""
import io
import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional dependency probes — missing libraries degrade gracefully
# ---------------------------------------------------------------------------
try:
    import pdfplumber  # noqa: F401

    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False
    logger.warning("pdfplumber not installed — PyPDF2 fallback will be used for PDFs")

try:
    import PyPDF2  # noqa: F401

    _HAS_PYPDF2 = True
except ImportError:
    _HAS_PYPDF2 = False

try:
    import pytesseract  # noqa: F401
    from PIL import Image  # noqa: F401

    _HAS_TESSERACT = True
except ImportError:
    _HAS_TESSERACT = False
    logger.warning(
        "pytesseract/Pillow not installed — image OCR disabled. "
        "Install tesseract-ocr and run: pip install pytesseract Pillow"
    )

try:
    from docx import Document as _DocxDocument  # noqa: F401

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logger.warning("python-docx not installed — DOCX extraction disabled")

# Maximum characters forwarded to the AI classifier (~2 000 tokens)
_MAX_TEXT_LENGTH = 8_000


class TextExtractor:
    """Dispatches text extraction to the appropriate backend by file extension."""

    def extract(self, filename: str, data: bytes) -> Optional[str]:
        """
        Extract plain text from *data* based on *filename*'s extension.

        Returns the extracted text (possibly truncated) or ``None`` on failure.
        """
        ext = os.path.splitext(filename.lower())[1]

        if ext == ".pdf":
            return self._extract_pdf(data)
        if ext in {".png", ".jpg", ".jpeg"}:
            return self._extract_image(data, filename)
        if ext == ".docx":
            return self._extract_docx(data, filename)

        logger.warning(f"No extractor registered for extension '{ext}' — skipping")
        return None

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _extract_pdf(self, data: bytes) -> Optional[str]:
        if _HAS_PDFPLUMBER:
            try:
                return self._pdf_via_pdfplumber(data)
            except Exception as exc:
                logger.warning(f"pdfplumber failed ({exc}) — trying PyPDF2 fallback")

        if _HAS_PYPDF2:
            try:
                return self._pdf_via_pypdf2(data)
            except Exception as exc:
                logger.error(f"PyPDF2 fallback also failed: {exc}")
                return None

        logger.error("No PDF extraction library available (install pdfplumber or PyPDF2)")
        return None

    def _pdf_via_pdfplumber(self, data: bytes) -> Optional[str]:
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)

        text = "\n".join(parts).strip()
        return text[:_MAX_TEXT_LENGTH] or None

    def _pdf_via_pypdf2(self, data: bytes) -> Optional[str]:
        import PyPDF2

        parts: list[str] = []
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)

        text = "\n".join(parts).strip()
        return text[:_MAX_TEXT_LENGTH] or None

    # ------------------------------------------------------------------
    # Images
    # ------------------------------------------------------------------

    def _extract_image(self, data: bytes, filename: str) -> Optional[str]:
        if not _HAS_TESSERACT:
            logger.warning(f"Skipping image '{filename}' — pytesseract not available")
            return None

        try:
            import pytesseract
            from PIL import Image

            image = Image.open(io.BytesIO(data))
            # Use German + English for invoice text recognition
            text = pytesseract.image_to_string(image, lang="deu+eng")
            text = text.strip()
            return text[:_MAX_TEXT_LENGTH] or None
        except Exception as exc:
            logger.error(f"OCR failed for '{filename}': {exc}")
            return None

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _extract_docx(self, data: bytes, filename: str) -> Optional[str]:
        if not _HAS_DOCX:
            logger.warning(f"Skipping DOCX '{filename}' — python-docx not available")
            return None

        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            parts: list[str] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())

            # Also pull text from tables (invoices often embed line items in tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())

            text = "\n".join(parts).strip()
            return text[:_MAX_TEXT_LENGTH] or None
        except Exception as exc:
            logger.error(f"DOCX extraction failed for '{filename}': {exc}")
            return None
